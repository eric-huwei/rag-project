from __future__ import annotations

import asyncio
from io import BytesIO
from pathlib import Path
import shutil
import unittest
import uuid
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
import fitz
from fastapi import UploadFile
from starlette.datastructures import Headers

from app.services.loading_service import (
    DOCX_CONTENT_TYPE,
    LoadingService,
    WORDPROCESSINGML_NS,
    get_chunking_config,
)


class LoadingServiceExtractionTests(unittest.TestCase):
    def setUp(self) -> None:
        self.test_root = Path(__file__).resolve().parent / ".tmp" / uuid.uuid4().hex
        self.test_root.mkdir(parents=True, exist_ok=True)
        self.service = LoadingService(root_dir=self.test_root)

    def tearDown(self) -> None:
        shutil.rmtree(self.test_root, ignore_errors=True)

    def _build_pdf_with_table(self) -> bytes:
        doc = fitz.open()
        page = doc.new_page(width=595, height=842)
        page.insert_text((50, 40), "Before table")

        x0, y0, cell_w, cell_h = 50, 100, 120, 32
        rows, cols = 3, 3
        for row in range(rows + 1):
            y = y0 + row * cell_h
            page.draw_line((x0, y), (x0 + cols * cell_w, y))
        for col in range(cols + 1):
            x = x0 + col * cell_w
            page.draw_line((x, y0), (x, y0 + rows * cell_h))

        values = [
            ["Name", "Qty", "Price"],
            ["Apple", "3", "9.9"],
            ["Banana", "5", "12.5"],
        ]
        for row, row_values in enumerate(values):
            for col, value in enumerate(row_values):
                rect = fitz.Rect(
                    x0 + col * cell_w + 4,
                    y0 + row * cell_h + 4,
                    x0 + (col + 1) * cell_w - 4,
                    y0 + (row + 1) * cell_h - 4,
                )
                page.insert_textbox(rect, value, fontsize=11)

        page.insert_text((50, 240), "After table")
        content = doc.tobytes()
        doc.close()
        return content

    def _build_pdf_with_image(self) -> bytes:
        doc = fitz.open()
        page = doc.new_page(width=300, height=300)
        page.insert_text((40, 40), "Before image")

        pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 40, 30), False)
        pix.clear_with(0x66CC33)
        page.insert_image(fitz.Rect(50, 80, 170, 150), pixmap=pix)

        page.insert_text((40, 220), "After image")
        content = doc.tobytes()
        doc.close()
        return content

    def _build_docx_with_table(self) -> bytes:
        document = Document()
        document.add_paragraph("第一段介绍。这里是第二句！")

        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Qty"
        table.cell(1, 0).text = "Apple"
        table.cell(1, 1).text = "3"

        document.add_paragraph("结尾段落，用于顺序验证。")

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _build_docx_for_sentence_split(self) -> bytes:
        document = Document()
        document.add_paragraph(
            "第一部分内容很长用于句子分块测试。第二部分内容继续用于句子分块测试！第三部分内容仍然用于句子分块测试？"
        )
        document.add_paragraph(
            "第四部分内容换段继续用于句子分块测试；第五部分内容换段继续用于句子分块测试。第六部分内容继续用于句子分块测试！"
        )
        document.add_paragraph(
            "第七部分内容在第三段里继续延展。第八部分内容在第三段里继续延展。第九部分内容在第三段里继续延展。"
        )

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _build_docx_with_auto_numbering(self) -> bytes:
        document = Document()
        document.add_paragraph("编号库1.", style="List Number")
        document.add_paragraph("编号库2.", style="List Number")
        document.add_paragraph("编号库1)", style="List Number 2")
        document.add_paragraph("编号库2)", style="List Number 2")

        buffer = BytesIO()
        document.save(buffer)
        return self._rewrite_docx_list_style(buffer.getvalue(), style_id="ListNumber2", lvl_text="%1)")

    def _rewrite_docx_list_style(self, content: bytes, *, style_id: str, lvl_text: str) -> bytes:
        buffer = BytesIO()
        with ZipFile(BytesIO(content)) as source, ZipFile(buffer, "w", compression=ZIP_DEFLATED) as target:
            numbering_root = ET.fromstring(source.read("word/numbering.xml"))
            for level in numbering_root.findall(f".//{{{WORDPROCESSINGML_NS}}}lvl"):
                pstyle = level.find(f"{{{WORDPROCESSINGML_NS}}}pStyle")
                if pstyle is None or pstyle.get(f"{{{WORDPROCESSINGML_NS}}}val") != style_id:
                    continue
                lvl_text_node = level.find(f"{{{WORDPROCESSINGML_NS}}}lvlText")
                if lvl_text_node is not None:
                    lvl_text_node.set(f"{{{WORDPROCESSINGML_NS}}}val", lvl_text)

            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "word/numbering.xml":
                    data = ET.tostring(numbering_root, encoding="utf-8", xml_declaration=True)
                target.writestr(item, data)
        return buffer.getvalue()

    def _make_upload_file(self, content: bytes, *, filename: str, content_type: str) -> UploadFile:
        return UploadFile(
            file=BytesIO(content),
            size=len(content),
            filename=filename,
            headers=Headers({"content-type": content_type}),
        )

    def test_extract_with_pymupdf_bytes_renders_table_as_markdown(self) -> None:
        page_map = self.service._extract_with_pymupdf_bytes(self._build_pdf_with_table())

        self.assertEqual(len(page_map), 1)
        page = page_map[0]
        text = page["text"]

        self.assertIn("Before table", text)
        self.assertIn("|Name|Qty|Price|", text)
        self.assertIn("|Apple|3|9.9|", text)
        self.assertIn("|Banana|5|12.5|", text)
        self.assertIn("After table", text)
        self.assertEqual(text.count("Apple"), 1)
        self.assertEqual(page["metadata"]["table_count"], 1)
        self.assertEqual(page["metadata"]["extraction_mode"], "pymupdf_table_aware")

    def test_clean_strategy_preserves_markdown_table_rows(self) -> None:
        page_map, method = self.service.load_pdf(
            self._build_pdf_with_table(),
            filename="table.pdf",
            loading_method="pymupdf",
            strategy="clean",
        )

        self.assertEqual(method, "pymupdf")
        self.assertIn("|---|---|---|", page_map[0]["text"])
        self.assertIn("Before table", page_map[0]["text"])
        self.assertIn("After table", page_map[0]["text"])

    def test_extract_with_pymupdf_bytes_tracks_images_and_positions(self) -> None:
        page_map = self.service._extract_with_pymupdf_bytes(self._build_pdf_with_image())

        self.assertEqual(len(page_map), 1)
        page = page_map[0]
        text = page["text"]

        self.assertIn("Before image", text)
        self.assertIn("[Image 1: format=png, size=40x30", text)
        self.assertIn("After image", text)
        self.assertEqual(page["metadata"]["image_count"], 1)
        self.assertEqual(page["metadata"]["extraction_mode"], "pymupdf_image_aware")
        self.assertEqual(page["metadata"]["images"][0]["ext"], "png")
        self.assertEqual(page["metadata"]["images"][0]["width"], 40)
        self.assertEqual(page["metadata"]["images"][0]["height"], 30)
        self.assertIn("_images", page)

    def test_persist_page_assets_writes_extracted_images(self) -> None:
        page_map = self.service._extract_with_pymupdf_bytes(self._build_pdf_with_image())

        summary = self.service._persist_page_assets(
            run_id="load_test_assets",
            filename="image.pdf",
            page_map=page_map,
        )

        self.assertEqual(summary["total_images"], 1)
        self.assertTrue(summary["asset_dir"])
        self.assertNotIn("_images", page_map[0])
        image_path = Path(page_map[0]["metadata"]["images"][0]["path"])
        self.assertTrue(image_path.exists())
        self.assertEqual(image_path.suffix.lower(), ".png")

    def test_load_docx_extracts_paragraphs_and_tables_in_body_order(self) -> None:
        upload = self._make_upload_file(
            self._build_docx_with_table(),
            filename="body-order.docx",
            content_type=DOCX_CONTENT_TYPE,
        )

        result = asyncio.run(
            self.service.load(
                upload,
                loading_method="auto",
            )
        )

        page_map = self.service.get_page_map()
        self.assertEqual(len(page_map), 1)
        page = page_map[0]
        text = page["text"]

        self.assertEqual(result["result"]["document_type"], "docx")
        self.assertEqual(result["result"]["loading_method"], "python-docx")
        self.assertEqual(result["loaded_content"]["metadata"]["document_type"], "docx")
        self.assertEqual(page["metadata"]["table_count"], 1)
        self.assertEqual(page["metadata"]["extraction_mode"], "python_docx_body_order")
        self.assertIn("第一段介绍。这里是第二句！", text)
        self.assertIn("|Name|Qty|", text)
        self.assertIn("|Apple|3|", text)
        self.assertIn("结尾段落，用于顺序验证。", text)
        self.assertLess(text.index("第一段介绍。这里是第二句！"), text.index("|Name|Qty|"))
        self.assertLess(text.index("|Apple|3|"), text.index("结尾段落，用于顺序验证。"))

    def test_docx_defaults_to_sentence_chunking_with_chinese_boundaries(self) -> None:
        upload = self._make_upload_file(
            self._build_docx_for_sentence_split(),
            filename="sentence-default.docx",
            content_type=DOCX_CONTENT_TYPE,
        )

        result = asyncio.run(
            self.service.load(
                upload,
                loading_method="auto",
                chunking_options={"max_chars": 100},
            )
        )

        page_map = self.service.get_page_map()
        chunks = result["loaded_content"]["chunks"]

        self.assertEqual(result["result"]["document_type"], "docx")
        self.assertEqual(result["result"]["chunking_strategy"], "sentence")
        self.assertEqual(result["loaded_content"]["metadata"]["chunking_strategy"], "sentence")
        self.assertEqual(result["loaded_content"]["metadata"]["total_pages"], 1)
        self.assertIn("\n\n", page_map[0]["text"])
        self.assertGreater(len(chunks), 1)
        self.assertTrue(any("！" in chunk["content"] for chunk in chunks))
        self.assertTrue(any("；" in chunk["content"] for chunk in chunks))

    def test_load_docx_restores_automatic_numbering_prefixes(self) -> None:
        page_map = self.service.load_docx(
            self._build_docx_with_auto_numbering(),
            filename="auto-numbering.docx",
            loading_method="python-docx",
        )

        text = page_map[0]["text"]
        self.assertIn("1. 编号库1.", text)
        self.assertIn("2. 编号库2.", text)
        self.assertIn("1) 编号库1)", text)
        self.assertIn("2) 编号库2)", text)

    def test_chunking_config_matches_backend_defaults(self) -> None:
        config = get_chunking_config()
        defaults = config["defaults_by_document_type"]
        strategies = {item["id"]: item for item in config["strategies"]}

        self.assertEqual(defaults["pdf"], "by_page")
        self.assertEqual(defaults["docx"], "sentence")
        self.assertIn("auto", strategies)
        self.assertIn("sentence", strategies)
        self.assertEqual(strategies["auto"]["request_value"], None)
        self.assertEqual(strategies["sentence"]["fields"][0]["key"], "max_chars")


if __name__ == "__main__":
    unittest.main()
